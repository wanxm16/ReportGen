"""Utility functions"""

import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from bs4 import BeautifulSoup
import re


def generate_file_id() -> str:
    """Generate unique file ID"""
    return str(uuid.uuid4())


def save_upload_file(file_content: bytes, filename: str, upload_dir: str) -> tuple[str, str]:
    """Save uploaded file and return file ID and path

    Args:
        file_content: File content bytes
        filename: Original filename
        upload_dir: Directory to save file

    Returns:
        Tuple of (file_id, file_path)
    """
    file_id = generate_file_id()
    file_extension = Path(filename).suffix
    new_filename = f"{file_id}{file_extension}"
    file_path = Path(upload_dir) / new_filename

    # Ensure directory exists
    Path(upload_dir).mkdir(parents=True, exist_ok=True)

    # Save file
    with open(file_path, 'wb') as f:
        f.write(file_content)

    return file_id, str(file_path)


def markdown_to_docx(markdown_text: str, output_path: str):
    """Convert markdown text to Word document

    Args:
        markdown_text: Markdown formatted text
        output_path: Path to save the Word document
    """
    doc = Document()

    # Parse markdown to HTML first
    html = markdown.markdown(markdown_text, extensions=['tables', 'nl2br'])
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.children:
        if element.name == 'h1':
            p = doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            p = doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            p = doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            p = doc.add_paragraph(element.get_text())
        elif element.name == 'table':
            # Extract table data
            rows = element.find_all('tr')
            if not rows:
                continue

            # Create table
            num_cols = len(rows[0].find_all(['th', 'td']))
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Light Grid Accent 1'

            # Fill table
            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    table.cell(i, j).text = cell.get_text().strip()

            doc.add_paragraph()  # Add spacing after table
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')

    doc.save(output_path)
